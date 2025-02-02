import re
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import latex2mathml.converter
import mathml2omml

from .utils import extract_latex_and_tag, create_math_xml


def process_run(run, para, doc):
    parts = re.split(r'(\$\$.+?\$\$)', run.text)
    if len(parts) == 1:
        return

    run_style = run.style
    run.text = parts[0]
    current_xml = run._r

    for part in parts[1:]:
        if part.startswith('$$') and part.endswith('$$'):
            latex_expr = part[2:-2].strip()
            latex_content, tag = extract_latex_and_tag(latex_expr)
            
            if tag:
                table = doc.add_table(rows=1, cols=2)
                table.allow_autofit = False
                table.columns[0].width = Inches(5)      # 80% of the page width
                table.columns[1].width = Inches(1.2)    # 20% for the tag

                cell_math = table.cell(0, 0)
                cell_math.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                math_para = cell_math.paragraphs[0]
                math_elem = create_math_xml(latex_content, inline=False)
                math_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                math_para._p.append(math_elem)


                cell_tag = table.cell(0, 1)
                cell_tag.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                tag_para = cell_tag.paragraphs[0]
                tag_elem = create_math_xml(f"({tag})", inline=False)
                tag_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                tag_para._p.append(tag_elem)

                current_xml.addnext(table._element)
                current_xml = table._element
            else:
                math_elem = create_math_xml(latex_expr, inline=True)
                current_xml.addnext(math_elem)
                current_xml = math_elem
        else:
            if part:
                new_run = para.add_run(part)
                new_run.style = run_style
                current_xml.addnext(new_run._r)
                current_xml = new_run._r


def process_paragraphs(doc):
    for para in list(doc.paragraphs):
        for run in list(para.runs):
            if '$$' in run.text:
                process_run(run, para, doc)


def convert_docx(input_path, output_path):
    doc = Document(input_path)
    process_paragraphs(doc)
    doc.save(output_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python -m mathmorph <input.docx> <output.docx>")
    else:
        convert_docx(sys.argv[1], sys.argv[2])
